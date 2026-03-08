"""
Report Export Module
===================
Export intelligence reports to various formats (Markdown, PDF, Word).
"""

import os
from datetime import datetime
from typing import Dict, List, Optional
import re

try:
    from fpdf import FPDF
except ImportError:
    FPDF = None

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None


class PDFReport(FPDF):
    def __init__(self):
        super().__init__()
    
    def header(self):
        self.set_font('Arial', 'B', 15)
        self.cell(0, 10, 'IntelNexus Intelligence Report', 0, 1, 'C')
        self.ln(5)
    
    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')


def export_markdown(content: str, query: str, output_path: str) -> str:
    """Export to Markdown format."""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(f"# IntelNexus Intelligence Report\n\n")
        f.write(f"**Query**: {query}\n")
        f.write(f"**Generated**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        f.write("---\n\n")
        f.write(content)
    return output_path


def export_pdf(content: str, query: str, output_path: str) -> str:
    """Export to PDF format."""
    if FPDF is None:
        raise ImportError("fpdf is not installed")
    
    pdf = PDFReport()
    pdf.add_page()
    
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, 'IntelNexus Intelligence Report', 0, 1, 'C')
    pdf.ln(10)
    
    pdf.set_font('Arial', '', 11)
    pdf.cell(0, 8, f'Query: {query}', 0, 1)
    pdf.cell(0, 8, f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', 0, 1)
    pdf.ln(10)
    
    pdf.set_fill_color(240, 240, 240)
    pdf.cell(0, 1, '', 0, 1, 'L', 1)
    pdf.ln(5)
    
    pdf.set_font('Arial', '', 10)
    pdf.multi_cell(0, 5, content[:5000] if len(content) > 5000 else content)
    
    pdf.output(output_path)
    return output_path


def export_word(content: str, query: str, output_path: str) -> str:
    """Export to Word format."""
    if Document is None:
        raise ImportError("python-docx is not installed")
    
    doc = Document()
    
    title = doc.add_heading('IntelNexus Intelligence Report', 0)
    
    doc.add_paragraph(f'Query: {query}')
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('')
    
    lines = content.split('\n')
    for line in lines:
        if not line.strip():
            continue
        
        if line.startswith('#'):
            level = min(line.count('#'), 3)
            text = line.replace('#', '').strip()
            doc.add_heading(text, level=level)
        elif line.startswith('**') and line.endswith('**'):
            text = line.replace('*', '').strip()
            p = doc.add_paragraph(text)
            if p.runs:
                p.runs[0].bold = True
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(line, style='List Number')
        else:
            if line.strip():
                doc.add_paragraph(line)
    
    doc.save(output_path)
    return output_path


def export_report(content: str, query: str, output_path: str, format: str = 'md') -> str:
    """Export report to specified format."""
    if not output_path:
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        output_path = f"report_{timestamp}"
    
    if format == 'pdf':
        if not output_path.endswith('.pdf'):
            output_path += '.pdf'
        return export_pdf(content, query, output_path)
    elif format == 'docx':
        if not output_path.endswith('.docx'):
            output_path += '.docx'
        return export_word(content, query, output_path)
    else:
        if not output_path.endswith('.md'):
            output_path += '.md'
        return export_markdown(content, query, output_path)


def get_export_formats() -> List[str]:
    """Get list of available export formats."""
    formats = ['md']
    if FPDF:
        formats.append('pdf')
    if Document:
        formats.append('docx')
    return formats
