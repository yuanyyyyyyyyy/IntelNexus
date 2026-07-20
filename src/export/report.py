"""
Report Export Module
===================
Export intelligence reports to various formats (Markdown, PDF, Word).
Supports Chinese and English with professional formatting.
"""

import os
from datetime import datetime
from typing import List, Optional
import re

from shared.logger import get_logger

logger = get_logger(__name__)

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    Workbook = None
    OPENPYXL_AVAILABLE = False


def _format_content_for_pdf(content: str) -> str:
    """Format content for better PDF rendering."""
    # 移除markdown的某些格式符号，使其在PDF中更清晰
    lines = content.split('\n')
    formatted_lines = []
    
    for line in lines:
        # 转换markdown标题
        if line.startswith('# '):
            formatted_lines.append('\n' + line.replace('# ', '■ ').upper())
        elif line.startswith('## '):
            formatted_lines.append('\n▸ ' + line.replace('## ', '').strip())
        else:
            formatted_lines.append(line)
    
    return '\n'.join(formatted_lines)


def export_markdown(content: str, query: str, output_path: str) -> str:
    """Export to Markdown format with enhanced structure."""
    # 清理内容，移除所有特殊字符
    content = _clean_content(content)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("# IntelNexus 智能情报报告\n\n")
        f.write(f"## 报告信息\n\n")
        f.write(f"- **查询内容**: {query}\n")
        f.write(f"- **生成时间**: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}\n")
        f.write(f"- **报告类型**: 多源网络情报分析\n\n")
        f.write("---\n\n")
        f.write("## 分析结果\n\n")
        f.write(content)
        f.write("\n\n---\n\n")
        f.write(f"*报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n")
        f.write("*© 2026 IntelNexus Platform - 多源网络情报分析平台*\n")
    return output_path


def _clean_markdown_for_word(text: str) -> str:
    """清理Markdown标记符号用于Word导出。"""
    # 移除markdown标题标记
    text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
    # 处理粗体：**text** -> text
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    # 处理斜体
    text = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', text)
    text = re.sub(r'(?<!_)_(?!_)(.+?)(?<!_)_(?!_)', r'\1', text)
    # 处理代码块
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'```[\s\S]*?```', '', text)
    # 处理链接
    text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'\1 (\2)', text)
    return text


def _clean_content(content: str) -> str:
    """清理内容特殊字符，用于所有导出格式。"""
    if not content:
        return content
    
    # 逐个替换特殊字符
    chars_to_remove = [
        '■', '□', '▢', '▣', '▤', '▥', '▦', '▧', '▨', '▩', '▪', '▫', '▬', '▭', '▮', '▯',
        '▰', '▱', '△', '▽', '▷', '◁', '◆', '◇', '○', '●', '◐', '◑', '◒', '◓', '◔', '◕',
        '◖', '◗', '★', '☆', '☉', '♠', '♣', '♥', '♦', '♩', '♪', '♫', '⚐', '⚑', '⚡',
        '⚪', '⚫', '⚬', '✓', '✗', '✘', '✔', '✖', '✚', '✽', '✿', '❀', '❖', '❤',
    ]
    for char in chars_to_remove:
        content = content.replace(char, '')
    
    # 移除emoji范围
    try:
        emoji_pattern = re.compile("["
            u"\U0001F600-\U0001F64F"
            u"\U0001F300-\U0001F5FF"
            u"\U0001F680-\U0001F6FF"
            u"\U0001F1E0-\U0001F1FF"
            "]+", flags=re.UNICODE)
        content = emoji_pattern.sub('', content)
    except Exception:
        pass
    
    return content


def _register_chinese_font():
    """Register Chinese font for PDF rendering."""
    if not REPORTLAB_AVAILABLE:
        return None
    font_paths = [
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/msyh.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/System/Library/Fonts/PingFang.ttc",
    ]
    for font_path in font_paths:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont("Chinese", font_path))
                return "Chinese"
            except Exception:
                continue
    return None


def _build_pdf_styles(font_name):
    """Build paragraph styles for PDF generation."""
    styles = getSampleStyleSheet()
    font_for_cjk = font_name if font_name else "Helvetica"

    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName=font_for_cjk,
        fontSize=20,
        textColor=colors.HexColor('#1F4E88'),
        spaceAfter=20,
        alignment=1
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontName=font_for_cjk,
        fontSize=14,
        textColor=colors.HexColor('#1F4E88'),
        spaceAfter=10,
        spaceBefore=15
    )
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName=font_for_cjk,
        fontSize=10,
        spaceAfter=8,
        alignment=0
    )
    sub_heading_style = ParagraphStyle(
        'SubHeading',
        parent=styles['Heading3'],
        fontName=font_for_cjk,
        fontSize=12,
        textColor=colors.HexColor('#2E5A88'),
        spaceAfter=8,
        spaceBefore=10
    )
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontName=font_for_cjk,
        fontSize=8,
        textColor=colors.grey,
        alignment=1
    )
    return {
        "title": title_style,
        "heading": heading_style,
        "normal": normal_style,
        "sub_heading": sub_heading_style,
        "footer": footer_style,
    }


def _content_to_pdf_story(content, styles):
    """Convert cleaned content to a list of PDF flowable elements."""
    story = []
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 5))
            continue
        if line.startswith('# '):
            clean_title = re.sub(r'^#+\s+', '', line)
            story.append(Paragraph(clean_title, styles["title"]))
        elif line.startswith('## '):
            clean_title = re.sub(r'^#+\s+', '', line)
            story.append(Paragraph(clean_title, styles["heading"]))
        elif line.startswith('### '):
            clean_title = re.sub(r'^#+\s+', '', line)
            story.append(Paragraph(clean_title, styles["sub_heading"]))
        elif line:
            story.append(Paragraph(line, styles["normal"]))
    return story


def export_pdf(content: str, query: str, output_path: str) -> str:
    """Export to PDF format with Chinese support using reportlab."""
    if not REPORTLAB_AVAILABLE:
        raise ImportError("reportlab is not installed. Install with: pip install reportlab")
    
    try:
        clean_query = query[:100] if query else "[No query content]"
    except Exception:
        clean_query = "[Query processing error]"
    
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    chinese_font = _register_chinese_font()
    styles = _build_pdf_styles(chinese_font)
    
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50
    )
    
    story = []
    
    # Title
    story.append(Paragraph("IntelNexus Intelligence Report", styles["title"]))
    story.append(Spacer(1, 10))
    story.append(Paragraph("<hr />", styles["normal"]))
    story.append(Spacer(1, 10))
    
    # Report Info
    story.append(Paragraph("Report Information", styles["heading"]))
    info_data = [
        ("Query:", clean_query),
        ("Generated:", datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ("Platform:", "IntelNexus v1.0"),
        ("Type:", "Multi-Source Network Intelligence Analysis")
    ]
    for label, value in info_data:
        story.append(Paragraph(f"<b>{label}</b> {value}", styles["normal"]))
    
    story.append(Spacer(1, 15))
    story.append(Paragraph("<hr />", styles["normal"]))
    story.append(Spacer(1, 10))
    
    # Analysis Results
    story.append(Paragraph("Analysis Results", styles["heading"]))
    
    max_length = 15000
    if len(content) > max_length:
        display_content = content[:max_length] + "\n\n[Content too long. Please check the full Markdown or Word report.]"
    else:
        display_content = content
    
    display_content = _clean_content(display_content)
    display_content = _clean_markdown_for_word(display_content)
    
    story.extend(_content_to_pdf_story(display_content, styles))
    
    # Footer
    story.append(Spacer(1, 20))
    story.append(Paragraph("<hr />", styles["normal"]))
    story.append(Paragraph(f"© 2026 IntelNexus Platform | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles["footer"]))
    
    doc.build(story)
    return output_path



def _add_paragraph_with_formatting(doc, text: str, style: str = None):
    """Add a paragraph to document with markdown formatting support.
    
    Converts **text** to bold, *text* to italic, `code` to code formatting.
    """
    if not text.strip():
        return
    
    para = doc.add_paragraph(style=style)
    
    # 粗体：**text**
    bold_pattern = r'\*\*(.+?)\*\*'
    # 斜体：*text*
    italic_pattern = r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)'
    # 代码：`text`
    code_pattern = r'`([^`]+)`'
    
    # 合并所有模式
    combined = f'({bold_pattern}|{italic_pattern}|{code_pattern})'
    
    last_end = 0
    for match in re.finditer(combined, text):
        # 添加之前的文本
        if match.start() > last_end:
            run = para.add_run(text[last_end:match.start()])
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        
        # 检查匹配的类型
        if match.group(2):  # 粗体
            run = para.add_run(match.group(2))
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        elif match.group(3):  # 斜体
            run = para.add_run(match.group(3))
            run.font.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        elif match.group(4):  # 代码
            run = para.add_run(match.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 0, 0)
        
        last_end = match.end()
    
    # 添加剩余的文本
    if last_end < len(text):
        run = para.add_run(text[last_end:])
        run.font.name = 'Calibri'
        run.font.size = Pt(11)


def export_word(content: str, query: str, output_path: str) -> str:
    """Export to Word format with markdown formatting rendering."""
    if Document is None:
        raise ImportError("python-docx is not installed. Install with: pip install python-docx")
    
    doc = Document()
    
    # 设置默认字体为支持中文的字体
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # 标题
    title = doc.add_heading('IntelNexus 智能情报分析报告', 0)
    title_format = title.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 报告信息
    info_heading = doc.add_heading('报告信息', level=1)
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_data = [
        ('查询内容', query if query else '[No query]'),
        ('生成时间', datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')),
        ('平台版本', 'IntelNexus v1.0'),
        ('报告类型', '多源网络情报分析')
    ]
    
    for i, (key, value) in enumerate(info_data):
        cells = info_table.rows[i].cells
        cells[0].text = key
        cells[1].text = str(value)
        # 设置格式
        for paragraph in cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    doc.add_paragraph()  # 空行
    
    # 分析结果
    result_heading = doc.add_heading('分析结果', level=1)
    
    # 清理内容，移除所有特殊字符
    content = _clean_content(content)
    
    # 处理markdown格式的内容 - 正确渲染markdown格式
    lines = content.split('\n')
    for line in lines:
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # 处理标题
        if line.startswith('# '):
            title_text = line.replace('# ', '').strip()
            heading = doc.add_heading(title_text, level=1)
        elif line.startswith('## '):
            title_text = line.replace('## ', '').strip()
            heading = doc.add_heading(title_text, level=2)
        elif line.startswith('### '):
            title_text = line.replace('### ', '').strip()
            heading = doc.add_heading(title_text, level=3)
        # 处理列表
        elif re.match(r'^\d+\.\s', line):
            list_text = re.sub(r'^\d+\.\s', '', line).strip()
            _add_paragraph_with_formatting(doc, list_text, 'List Number')
        elif line.startswith('- '):
            list_text = line[2:].strip()
            _add_paragraph_with_formatting(doc, list_text, 'List Bullet')
        elif line.startswith('* '):
            list_text = line[2:].strip()
            _add_paragraph_with_formatting(doc, list_text, 'List Bullet')
        else:
            # 清理可能的markdown标题标记（处理行内或意外的情况）
            cleaned_line = _clean_markdown_for_word(line.strip())
            if cleaned_line.strip():
                _add_paragraph_with_formatting(doc, cleaned_line)
    
    # 添加页脚
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"© 2026 IntelNexus Platform | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    doc.save(output_path)
    return output_path



def export_report(content: str, query: str, output_path: str, format: str = 'md') -> str:
    """Export report to specified format."""
    if not output_path:
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        output_path = f"report_{timestamp}"
    
    if format == 'pdf':
        if not output_path.endswith('.pdf'):
            output_path += '.pdf'
        return export_pdf(content, query, output_path)
    elif format == 'docx':
        if not output_path.endswith('.docx'):
            output_path += '.docx'
        return export_word(content, query, output_path)
    else:
        if not output_path.endswith('.md'):
            output_path += '.md'
        return export_markdown(content, query, output_path)


def get_export_formats() -> List[str]:
    """Get list of available export formats."""
    formats = ['md']
    if REPORTLAB_AVAILABLE:
        formats.append('pdf')
    if Document:
        formats.append('docx')
    if OPENPYXL_AVAILABLE:
        formats.append('xlsx')
    return formats


def export_excel(content: str, query: str, output_path: str) -> str:
    """Export to Excel format with proper formatting."""
    if Workbook is None:
        raise ImportError("openpyxl is not installed. Install with: pip install openpyxl")
    
    wb = Workbook()
    ws = wb.active
    ws.title = "情报报告"
    
    # 定义样式
    header_font = Font(name='微软雅黑', size=16, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center')
    
    title_font = Font(name='微软雅黑', size=12, bold=True)
    title_fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
    
    normal_font = Font(name='微软雅黑', size=11)
    wrap_alignment = Alignment(wrap_text=True, vertical='top')
    
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # 标题行
    ws.merge_cells('A1:B1')
    ws['A1'] = 'IntelNexus 智能情报分析报告'
    ws['A1'].font = header_font
    ws['A1'].fill = header_fill
    ws['A1'].alignment = header_alignment
    ws.row_dimensions[1].height = 30
    
    # 报告信息
    ws['A3'] = '查询内容'
    ws['B3'] = query if query else '[无查询内容]'
    ws['A4'] = '生成时间'
    ws['B4'] = datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')
    ws['A5'] = '平台版本'
    ws['B5'] = 'IntelNexus v1.0'
    ws['A6'] = '报告类型'
    ws['B6'] = '多源网络情报分析'
    
    for row in range(3, 7):
        ws[f'A{row}'].font = title_font
        ws[f'A{row}'].fill = title_fill
        ws[f'A{row}'].border = thin_border
        ws[f'B{row}'].border = thin_border
        ws[f'B{row}'].alignment = wrap_alignment
    
    # 分析结果标题
    ws['A8'] = '分析结果'
    ws['A8'].font = title_font
    ws['A8'].fill = title_fill
    ws.merge_cells('A8:B8')
    ws['A8'].alignment = Alignment(horizontal='center', vertical='center')
    ws['A8'].border = thin_border
    ws['B8'].border = thin_border
    ws.row_dimensions[8].height = 25
    
    # 解析内容并添加到 Excel
    start_row = 9
    current_row = start_row
    
    # 清理内容中的markdown标题标记
    clean_content = _clean_markdown_for_word(content)
    
    # 按段落添加内容
    paragraphs = clean_content.split('\n\n')
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        # 检查是否是标题
        is_title = False
        if para.startswith('■ ') or para.startswith('▸ '):
            is_title = True
            para = para[2:].strip() if para.startswith('■ ') else para[2:].strip()
        
        ws[f'A{current_row}'] = para
        ws.merge_cells(f'A{current_row}:B{current_row}')
        
        if is_title:
            ws[f'A{current_row}'].font = Font(name='微软雅黑', size=11, bold=True, color='1F4E79')
        else:
            ws[f'A{current_row}'].font = normal_font
        
        ws[f'A{current_row}'].alignment = wrap_alignment
        ws[f'A{current_row}'].border = thin_border
        ws.row_dimensions[current_row].height = max(20, len(para) // 40 * 15 + 20)
        
        current_row += 1
    
    # 设置列宽
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 70
    
    # 添加页脚
    footer_row = current_row + 2
    ws.merge_cells(f'A{footer_row}:B{footer_row}')
    ws[f'A{footer_row}'] = f"© 2026 IntelNexus Platform | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    ws[f'A{footer_row}'].font = Font(name='微软雅黑', size=9, italic=True, color='808080')
    ws[f'A{footer_row}'].alignment = Alignment(horizontal='center')
    
    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    if not output_path.endswith('.xlsx'):
        output_path += '.xlsx'
    
    wb.save(output_path)
    return output_path
