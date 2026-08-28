"""
Report Export Module
===================
Export intelligence reports to various formats (Markdown, PDF, Word).
Supports Chinese and English with professional formatting.
"""

import os
from datetime import datetime
from typing import List, Optional
import re

from intelnexus.core.logger import get_logger
from intelnexus.export.font_registry import DOCX_CJK_FONT_NAME

logger = get_logger(__name__)

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    Workbook = None
    OPENPYXL_AVAILABLE = False


def _format_content_for_pdf(content: str) -> str:
    """Format content for better PDF rendering."""
    # 移除markdown的某些格式符号，使其在PDF中更清晰
    lines = content.split('\n')
    formatted_lines = []
    
    for line in lines:
        # 转换markdown标题
        if line.startswith('# '):
            formatted_lines.append('\n' + line.replace('# ', '■ ').upper())
        elif line.startswith('## '):
            formatted_lines.append('\n▸ ' + line.replace('## ', '').strip())
        else:
            formatted_lines.append(line)
    
    return '\n'.join(formatted_lines)


def export_markdown(content: str, query: str, output_path: str) -> str:
    """Export to Markdown format.

    如果 content 已经是 10 板块结构化报告（以 "# IntelNexus" 或 "=" 开头），直接写入；
    否则回退到旧版包装格式（向后兼容）。
    """
    is_structured = (content.startswith("# IntelNexus") or
                     content.startswith("=") or
                     "## 二、核心摘要" in content[:500] or
                     "## 三、事件画像" in content[:1000])

    with open(output_path, 'w', encoding='utf-8') as f:
        if is_structured:
            # 新版 10 板块结构化报告，直接写入
            f.write(content)
        else:
            # 旧版 LLM 原始输出，包装为简单报告
            f.write("# IntelNexus 智能情报报告\n\n")
            f.write(f"## 报告信息\n\n")
            f.write(f"- **查询内容**: {query}\n")
            f.write(f"- **生成时间**: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}\n")
            f.write(f"- **报告类型**: 多源网络情报分析\n\n")
            f.write("---\n\n")
            f.write("## 分析结果\n\n")
            f.write(_clean_content(content))
            f.write("\n\n---\n\n")
            f.write(f"*报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n")
            f.write("*© 2026 IntelNexus Platform - 多源网络情报分析平台*\n")
    return output_path


def _clean_markdown_for_word(text: str) -> str:
    """清理Markdown标记符号用于Word导出。"""
    # 移除markdown标题标记
    text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
    # 处理粗体：**text** -> text
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    # 处理斜体
    text = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', text)
    text = re.sub(r'(?<!_)_(?!_)(.+?)(?<!_)_(?!_)', r'\1', text)
    # 处理代码块
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'```[\s\S]*?```', '', text)
    # 处理链接
    text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'\1 (\2)', text)
    # 剥离证据角标与 HTML 残留：<sup>[N]</sup>、<p ...>...</p>、<img ...>、<b>/<br>
    text = re.sub(r'<sup>\[\d+\]</sup>', '', text)
    text = re.sub(r'</?sup>', '', text)
    text = re.sub(r'<img[^>]*>', '', text)
    text = re.sub(r'<p[^>]*>|</p>', '', text)
    text = re.sub(r'<br\s*/?>', chr(10), text)
    return text


def _clean_content(content: str) -> str:
    """清理内容特殊字符，用于所有导出格式。"""
    if not content:
        return content
    
    # 逐个替换特殊字符
    chars_to_remove = [
        '■', '□', '▢', '▣', '▤', '▥', '▦', '▧', '▨', '▩', '▪', '▫', '▬', '▭', '▮', '▯',
        '▰', '▱', '△', '▽', '▷', '◁', '◆', '◇', '○', '●', '◐', '◑', '◒', '◓', '◔', '◕',
        '◖', '◗', '★', '☆', '☉', '♠', '♣', '♥', '♦', '♩', '♪', '♫', '⚐', '⚑', '⚡',
        '⚪', '⚫', '⚬', '✓', '✗', '✘', '✔', '✖', '✚', '✽', '✿', '❀', '❖', '❤',
    ]
    for char in chars_to_remove:
        content = content.replace(char, '')
    
    # 移除emoji范围
    try:
        emoji_pattern = re.compile("["
            u"\U0001F600-\U0001F64F"
            u"\U0001F300-\U0001F5FF"
            u"\U0001F680-\U0001F6FF"
            u"\U0001F1E0-\U0001F1FF"
            "]+", flags=re.UNICODE)
        content = emoji_pattern.sub('', content)
    except Exception:
        pass
    
    return content




# --------------------------------------------------------------------------
# Markdown 表格解析（关键数据等板块在三种导出中渲染为真实表格）
# --------------------------------------------------------------------------
def _parse_md_tables(lines: List[str]) -> List[dict]:
    """把连续的 markdown 表格行解析为 [{header: [...], rows: [[...]]}]。"""
    tables = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:\-|]+\|$", lines[i + 1].strip()):
            header = [c.strip() for c in line.strip("|").split("|")]
            i += 2  # 跳过分隔行
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            tables.append({"header": header, "rows": rows})
        else:
            i += 1
    return tables

def _register_chinese_font():
    """Register Chinese font for PDF rendering.

    统一走 font_registry：优先项目自带 Noto Sans SC（Regular + Bold 真字重），
    缺失时按系统字体候选链兜底。保留既有注册名 ``Chinese``。

    Returns:
        tuple: ``(regular_name | None, bold_name | None)``
    """
    if not REPORTLAB_AVAILABLE:
        return None, None
    from intelnexus.export.font_registry import register_pdf_fonts

    return register_pdf_fonts("Chinese", "Chinese-Bold")


def _build_pdf_styles(font_name, bold_name=None):
    """Build paragraph styles for PDF generation.

    标题/章节标题使用 Bold 真字重，正文使用 Regular；
    段落内 ``<b>`` 经 registerFontFamily 映射到 Bold。
    """
    styles = getSampleStyleSheet()
    font_for_cjk = font_name if font_name else "Helvetica"
    font_bold_cjk = bold_name if bold_name else font_for_cjk

    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName=font_bold_cjk,
        fontSize=20,
        textColor=colors.HexColor('#1F4E88'),
        spaceAfter=20,
        alignment=1
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontName=font_bold_cjk,
        fontSize=14,
        textColor=colors.HexColor('#1F4E88'),
        spaceAfter=10,
        spaceBefore=15
    )
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName=font_for_cjk,
        fontSize=10,
        spaceAfter=8,
        alignment=0
    )
    sub_heading_style = ParagraphStyle(
        'SubHeading',
        parent=styles['Heading3'],
        fontName=font_bold_cjk,
        fontSize=12,
        textColor=colors.HexColor('#2E5A88'),
        spaceAfter=8,
        spaceBefore=10
    )
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontName=font_for_cjk,
        fontSize=8,
        textColor=colors.grey,
        alignment=1
    )
    return {
        "title": title_style,
        "heading": heading_style,
        "normal": normal_style,
        "sub_heading": sub_heading_style,
        "footer": footer_style,
    }


def _content_to_pdf_story(content, styles):
    """Convert cleaned content to a list of PDF flowable elements.

    每行先做 XML 转义再进 Paragraph：LLM 报告可能含 <script>、未闭合标签、
    裸 & 等字符，paraparser 会当作 XML 解析导致 "unclosed tags" 崩溃。
    markdown 表格块（如「五、关键数据」）渲染为真正的 reportlab Table。
    """
    from xml.sax.saxutils import escape as _xml_escape
    # 剥离导出中无意义的 HTML 残留（角标/图表/过渡段落），否则被 escape 成字面量刷屏
    content = re.sub(r"<sup>\[\d+\]</sup>", "", content)
    content = re.sub(r"</?sup>", "", content)
    content = re.sub(r"<img[^>]*>", "", content)
    content = re.sub(r"<p[^>]*>|</p>", "", content)
    story = []
    lines = content.split("\n")

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        # 表格块：header 行 + 分隔行 + 数据行 -> reportlab Table
        if (stripped.startswith("|") and i + 1 < len(lines)
                and re.match(r"^\|[\s:\-|]+\|$", lines[i + 1].strip())):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table_data = [[_xml_escape(c) for c in header]]
            for row in rows:
                table_data.append([_xml_escape(c) for c in row])
            tbl = Table(table_data)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E1F2")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F4E88")),
                ("FONTNAME", (0, 0), (-1, -1), styles["normal"].fontName),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(Spacer(1, 4))
            story.append(tbl)
            story.append(Spacer(1, 6))
            continue

        if not stripped:
            story.append(Spacer(1, 5))
        elif stripped.startswith("### "):
            story.append(Paragraph(_xml_escape(re.sub(r"^#+\s+", "", stripped)), styles["sub_heading"]))
        elif stripped.startswith("## "):
            story.append(Paragraph(_xml_escape(re.sub(r"^#+\s+", "", stripped)), styles["heading"]))
        elif stripped.startswith("# "):
            story.append(Paragraph(_xml_escape(re.sub(r"^#+\s+", "", stripped)), styles["title"]))
        else:
            body = _xml_escape(stripped)
            body = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", body)
            story.append(Paragraph(body, styles["normal"]))
        i += 1
    return story


def export_pdf(content: str, query: str, output_path: str) -> str:
    """Export to PDF format with Chinese support using reportlab."""
    if not REPORTLAB_AVAILABLE:
        raise ImportError("reportlab is not installed. Install with: pip install reportlab")
    
    try:
        clean_query = query[:100] if query else "[No query content]"
    except Exception:
        clean_query = "[Query processing error]"
    
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    chinese_font, chinese_bold = _register_chinese_font()
    styles = _build_pdf_styles(chinese_font, chinese_bold)
    
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50
    )
    
    story = []
    
    is_structured = (content.startswith("# IntelNexus") or
                     content.startswith("=") or
                     "## 二、核心摘要" in content[:500] or
                     "## 三、事件画像" in content[:1000])
    
    if not is_structured:
        # 旧版格式：添加标题和报告信息
        story.append(Paragraph("IntelNexus Intelligence Report", styles["title"]))
        story.append(Spacer(1, 10))
        story.append(Paragraph("<hr />", styles["normal"]))
        story.append(Spacer(1, 10))
        
        story.append(Paragraph("Report Information", styles["heading"]))
        info_data = [
            ("Query:", clean_query),
            ("Generated:", datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ("Platform:", "IntelNexus v1.0"),
            ("Type:", "Multi-Source Network Intelligence Analysis")
        ]
        for label, value in info_data:
            story.append(Paragraph(f"<b>{label}</b> {value}", styles["normal"]))
        
        story.append(Spacer(1, 15))
        story.append(Paragraph("<hr />", styles["normal"]))
        story.append(Spacer(1, 10))
        
        story.append(Paragraph("Analysis Results", styles["heading"]))
    # 新版结构化报告：内容已包含完整标题和信息，直接渲染
    
    max_length = 15000
    if len(content) > max_length:
        display_content = content[:max_length] + "\n\n[Content too long. Please check the full Markdown or Word report.]"
    else:
        display_content = content
    
    display_content = _clean_content(display_content)
    display_content = _clean_markdown_for_word(display_content)
    
    story.extend(_content_to_pdf_story(display_content, styles))
    
    # Footer
    story.append(Spacer(1, 20))
    story.append(Paragraph("<hr />", styles["normal"]))
    story.append(Paragraph(f"© 2026 IntelNexus Platform | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles["footer"]))
    
    doc.build(story)
    return output_path



def _add_paragraph_with_formatting(doc, text: str, style: str = None):
    """Add a paragraph to document with markdown formatting support.
    
    Converts **text** to bold, *text* to italic, `code` to code formatting.
    """
    if not text.strip():
        return
    
    para = doc.add_paragraph(style=style)
    
    # 粗体：**text**
    bold_pattern = r'\*\*(.+?)\*\*'
    # 斜体：*text*
    italic_pattern = r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)'
    # 代码：`text`
    code_pattern = r'`([^`]+)`'
    
    # 合并所有模式
    combined = f'({bold_pattern}|{italic_pattern}|{code_pattern})'
    
    last_end = 0
    for match in re.finditer(combined, text):
        # 添加之前的文本
        if match.start() > last_end:
            run = para.add_run(text[last_end:match.start()])
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        
        # 检查匹配的类型
        if match.group(2):  # 粗体
            run = para.add_run(match.group(2))
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        elif match.group(3):  # 斜体
            run = para.add_run(match.group(3))
            run.font.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        elif match.group(4):  # 代码
            run = para.add_run(match.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 0, 0)
        
        last_end = match.end()
    
    # 添加剩余的文本
    if last_end < len(text):
        run = para.add_run(text[last_end:])
        run.font.name = 'Calibri'
        run.font.size = Pt(11)


def export_word(content: str, query: str, output_path: str) -> str:
    """Export to Word format with markdown formatting rendering."""
    if Document is None:
        raise ImportError("python-docx is not installed. Install with: pip install python-docx")
    
    doc = Document()
    
    # 设置默认字体：中英文都指定（w:eastAsia 缺失时 Word 用默认东亚字体渲染，
    # 部分环境显示为乱码/豆腐块）。中文用思源黑体（Source Han Sans SC）：
    # docx 仅按名称引用字体、不内嵌字体文件，无字体授权风险；接收方缺少该字体时，
    # Word 会按主题自动替换为可用中文字体，仍可正常渲染。
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    try:
        from docx.oxml.ns import qn
        style.element.rPr.rFonts.set(qn('w:eastAsia'), DOCX_CJK_FONT_NAME)
        for hname, hsize in (('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 12)):
            hs = doc.styles[hname]
            hs.font.name = 'Calibri'
            if hs.element.rPr is not None and hs.element.rPr.rFonts is not None:
                hs.element.rPr.rFonts.set(qn('w:eastAsia'), DOCX_CJK_FONT_NAME)
    except Exception as e:
        logger.debug(f"Word east-asian font setup skipped: {e}")
    
    # 标题 + 报告信息（仅旧版格式需要，新版结构化报告已包含）
    is_structured = (content.startswith("# IntelNexus") or
                     content.startswith("=") or
                     "## 二、核心摘要" in content[:500] or
                     "## 三、事件画像" in content[:1000])
    
    if not is_structured:
        title = doc.add_heading('IntelNexus 智能情报分析报告', 0)
        title_format = title.paragraph_format
        title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info_heading = doc.add_heading('报告信息', level=1)
        
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('查询内容', query if query else '[No query]'),
            ('生成时间', datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')),
            ('平台版本', 'IntelNexus v1.0'),
            ('报告类型', '多源网络情报分析')
        ]
        
        for i, (key, value) in enumerate(info_data):
            cells = info_table.rows[i].cells
            cells[0].text = key
            cells[1].text = str(value)
            for paragraph in cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        doc.add_paragraph()  # 空行
        
        result_heading = doc.add_heading('分析结果', level=1)
    
    # 清理内容，移除所有特殊字符
    content = _clean_content(content)
    
    # 处理markdown格式的内容 - 正确渲染markdown格式
    lines = content.split('\n')
    # markdown 表格块 -> 真实 Word 表格（关键数据板块）
    _md_tables = _parse_md_tables(lines)
    _tbl_iter = iter(_md_tables)
    _skip_pipe = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and not _skip_pipe:
            # 表格块开始：渲染真实 Word 表格并跳过全部管道行
            try:
                _t = next(_tbl_iter)
            except StopIteration:
                _skip_pipe = True
                continue
            _cols = len(_t["header"])
            _tbl = doc.add_table(rows=1 + len(_t["rows"]), cols=_cols)
            _tbl.style = 'Light Grid Accent 1'
            for _ci, _cv in enumerate(_t["header"]):
                _cell_p = _tbl.rows[0].cells[_ci].paragraphs[0]
                _run = _cell_p.add_run(_cv)
                _run.font.bold = True
            for _ri, _row in enumerate(_t["rows"], start=1):
                for _ci, _cv in enumerate(_row[:_cols]):
                    _tbl.rows[_ri].cells[_ci].text = _cv
            doc.add_paragraph()
            _skip_pipe = True
            continue
        if stripped.startswith("|") and _skip_pipe:
            # 同一表格块的剩余管道行：跳过（表格已整体渲染）
            continue
        _skip_pipe = False
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # 处理标题
        if line.startswith('# '):
            title_text = line.replace('# ', '').strip()
            heading = doc.add_heading(title_text, level=1)
        elif line.startswith('## '):
            title_text = line.replace('## ', '').strip()
            heading = doc.add_heading(title_text, level=2)
        elif line.startswith('### '):
            title_text = line.replace('### ', '').strip()
            heading = doc.add_heading(title_text, level=3)
        # 处理列表
        elif re.match(r'^\d+\.\s', line):
            list_text = re.sub(r'^\d+\.\s', '', line).strip()
            _add_paragraph_with_formatting(doc, list_text, 'List Number')
        elif line.startswith('- '):
            list_text = line[2:].strip()
            _add_paragraph_with_formatting(doc, list_text, 'List Bullet')
        elif line.startswith('* '):
            list_text = line[2:].strip()
            _add_paragraph_with_formatting(doc, list_text, 'List Bullet')
        else:
            # 清理可能的markdown标题标记（处理行内或意外的情况）
            cleaned_line = _clean_markdown_for_word(line.strip())
            if cleaned_line.strip():
                _add_paragraph_with_formatting(doc, cleaned_line)
    
    # 添加页脚
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"© 2026 IntelNexus Platform | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    doc.save(output_path)
    return output_path



def export_report(content: str, query: str, output_path: str, format: str = 'md') -> str:
    """Export report to specified format."""
    if not output_path:
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        output_path = f"report_{timestamp}"
    
    if format == 'pdf':
        if not output_path.endswith('.pdf'):
            output_path += '.pdf'
        return export_pdf(content, query, output_path)
    elif format == 'docx':
        if not output_path.endswith('.docx'):
            output_path += '.docx'
        return export_word(content, query, output_path)
    else:
        if not output_path.endswith('.md'):
            output_path += '.md'
        return export_markdown(content, query, output_path)


def get_export_formats() -> List[str]:
    """Get list of available export formats."""
    formats = ['md']
    if REPORTLAB_AVAILABLE:
        formats.append('pdf')
    if Document:
        formats.append('docx')
    if OPENPYXL_AVAILABLE:
        formats.append('xlsx')
    return formats


def export_excel(content: str, query: str, output_path: str) -> str:
    """Export to Excel format with proper formatting."""
    if Workbook is None:
        raise ImportError("openpyxl is not installed. Install with: pip install openpyxl")
    
    wb = Workbook()
    ws = wb.active
    ws.title = "情报报告"
    
    # 定义样式
    header_font = Font(name=DOCX_CJK_FONT_NAME, size=16, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center')
    
    title_font = Font(name=DOCX_CJK_FONT_NAME, size=12, bold=True)
    title_fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
    
    normal_font = Font(name=DOCX_CJK_FONT_NAME, size=11)
    wrap_alignment = Alignment(wrap_text=True, vertical='top')
    
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # 标题行 + 报告信息（仅旧版格式需要）
    is_structured = (content.startswith("# IntelNexus") or
                     content.startswith("=") or
                     "## 二、核心摘要" in content[:500] or
                     "## 三、事件画像" in content[:1000])
    start_row = 1
    
    if not is_structured:
        ws.merge_cells('A1:B1')
        ws['A1'] = 'IntelNexus 智能情报分析报告'
        ws['A1'].font = header_font
        ws['A1'].fill = header_fill
        ws['A1'].alignment = header_alignment
        ws.row_dimensions[1].height = 30
        
        ws['A3'] = '查询内容'
        ws['B3'] = query if query else '[无查询内容]'
        ws['A4'] = '生成时间'
        ws['B4'] = datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')
        ws['A5'] = '平台版本'
        ws['B5'] = 'IntelNexus v1.0'
        ws['A6'] = '报告类型'
        ws['B6'] = '多源网络情报分析'
        
        for row in range(3, 7):
            ws[f'A{row}'].font = title_font
            ws[f'A{row}'].fill = title_fill
            ws[f'A{row}'].border = thin_border
            ws[f'B{row}'].border = thin_border
            ws[f'B{row}'].alignment = wrap_alignment
        
        ws['A8'] = '分析结果'
        ws['A8'].font = title_font
        ws['A8'].fill = title_fill
        ws.merge_cells('A8:B8')
        ws['A8'].alignment = Alignment(horizontal='center', vertical='center')
        ws['A8'].border = thin_border
        ws['B8'].border = thin_border
        ws.row_dimensions[8].height = 25
        
        start_row = 9
    else:
        # 新版结构化报告：从第一行开始
        start_row = 1
    
    # 解析内容并添加到 Excel
    current_row = start_row
    
    # 清理内容中的markdown标题标记
    clean_content = _clean_markdown_for_word(content)
    # markdown 表格 -> 每行数据变为「表头: 值」拼接的可读文本（Excel 无表格对象）
    try:
        _xl_lines = clean_content.split(chr(10))
        for _t in _parse_md_tables(_xl_lines):
            _rendered = [
                chr(0xFF1B).join(f"{h}: {v}" for h, v in zip(_t['header'], row))
                for row in _t['rows']
            ]
            _first = next(i for i, l in enumerate(_xl_lines) if l.strip().startswith('|'))
            _block = chr(10).join(_xl_lines[_first:_first + 2 + len(_t['rows'])])
            clean_content = clean_content.replace(_block, chr(10).join(_rendered))
    except Exception as e:
        logger.debug(f'excel table flatten skipped: {e}')
    
    # 按段落添加内容
    paragraphs = clean_content.split('\n\n')
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        # 检查是否是标题
        is_title = False
        if para.startswith('■ ') or para.startswith('▸ '):
            is_title = True
            para = para[2:].strip() if para.startswith('■ ') else para[2:].strip()
        
        ws[f'A{current_row}'] = para
        ws.merge_cells(f'A{current_row}:B{current_row}')
        
        if is_title:
            ws[f'A{current_row}'].font = Font(name=DOCX_CJK_FONT_NAME, size=11, bold=True, color='1F4E79')
        else:
            ws[f'A{current_row}'].font = normal_font
        
        ws[f'A{current_row}'].alignment = wrap_alignment
        ws[f'A{current_row}'].border = thin_border
        ws.row_dimensions[current_row].height = max(20, len(para) // 40 * 15 + 20)
        
        current_row += 1
    
    # 设置列宽
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 70
    
    # 添加页脚
    footer_row = current_row + 2
    ws.merge_cells(f'A{footer_row}:B{footer_row}')
    ws[f'A{footer_row}'] = f"© 2026 IntelNexus Platform | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    ws[f'A{footer_row}'].font = Font(name=DOCX_CJK_FONT_NAME, size=9, italic=True, color='808080')
    ws[f'A{footer_row}'].alignment = Alignment(horizontal='center')
    
    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    if not output_path.endswith('.xlsx'):
        output_path += '.xlsx'
    
    wb.save(output_path)
    return output_path
